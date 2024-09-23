from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph

def convert_to_pdf_from_text(text):
    """
    Convert text to a PDF with styled headings and body text.
    
    Args:
        text: A string containing the text to be converted.
        
    Returns:
        Byte content of the generated PDF.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []

    # Define styles
    heading_style = ParagraphStyle(
        name='HeadingStyle',
        fontName='Helvetica-Bold',
        fontSize=14,
        spaceAfter=9,  # Space after heading
    )
    
    subheading_style = ParagraphStyle(
        name='SubheadingStyle',
        fontName='Helvetica-Bold',
        fontSize=12,
        spaceAfter=5,  # Space after subheading
    )
    
    body_style = ParagraphStyle(
        name='BodyStyle',
        fontName='Helvetica',
        fontSize=10,
        spaceAfter=2,  # Space after body text
    )

    # Process the text and apply styles
    paragraphs = text.split('\n')
    for para in paragraphs:
        para = para.strip()
        if para.startswith('# '):  # Heading
            styled_paragraph = Paragraph(para[2:], heading_style)
        elif para.startswith('## '):  # Subheading
            styled_paragraph = Paragraph(para[3:], subheading_style)
        else:  # Body text
            styled_paragraph = Paragraph(para, body_style)
        
        story.append(styled_paragraph)

    # Build the PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def convert_to_docx_from_text(text):
    """
    Convert text to a DOCX file with styled headings and body text.
    
    Args:
        text: A string containing the text to be converted.
        
    Returns:
        Byte content of the generated DOCX file.
    """
    buffer = BytesIO()
    doc = Document()
    
    # Define and apply styles
    heading_style = doc.styles['Heading1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_after = Pt(12)
    
    subheading_style = doc.styles['Heading2']
    subheading_style.font.name = 'Arial'
    subheading_style.font.size = Pt(12)
    subheading_style.font.bold = True
    subheading_style.paragraph_format.space_after = Pt(12)
    
    body_style = doc.styles['Normal']
    body_style.font.name = 'Arial'
    body_style.font.size = Pt(10)
    body_style.paragraph_format.line_spacing = Pt(10)
    body_style.paragraph_format.space_after = Pt(4)
    
    # Process the text and apply styles
    paragraphs = text.split('\n')
    for para in paragraphs:
        para = para.strip()
        if para.startswith('# '):  # Heading
            doc.add_paragraph(para[2:], style='Heading1')
        elif para.startswith('## '):  # Subheading
            doc.add_paragraph(para[3:], style='Heading2')
        else:  # Body text
            doc.add_paragraph(para, style='Normal')
    
    # Save the document to the BytesIO buffer
    doc.save(buffer)
    buffer.seek(0)  # Move the cursor to the beginning of the buffer
    return buffer.getvalue()  # Return the byte content of the DOCX file